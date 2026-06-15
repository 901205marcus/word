from __future__ import annotations

import re
from calendar import monthrange
from copy import deepcopy
from dataclasses import dataclass
from datetime import date, datetime, timedelta
from pathlib import Path

from docx import Document

from .models import ActionType, ApplySummary, ScheduleAction
from .parser import normalize_date, normalize_time, time_sort_key

DAY_HEADER_PATTERN = re.compile(r"(\d{1,2})/(\d{1,2})(?:\s*[\(（]|$)")
HEADER_WITH_WEEKDAY_PATTERN = re.compile(
    r"(?P<date>\d{1,2}/\d{1,2})(?P<space>\s*)(?P<open>[\(（])(?P<weekday>[一二三四五六日天])(?P<close>[\)）])"
)
WEEKDAY_MAP = "一二三四五六日"


@dataclass(slots=True)
class DayHeader:
    row_idx: int
    month: int
    day: int
    raw_text: str

    @property
    def date_text(self) -> str:
        return f"{self.month}/{self.day}"


def cell_text(cell) -> str:
    return "\n".join(paragraph.text.strip() for paragraph in cell.paragraphs).strip()


def row_texts(table, row_idx: int) -> list[str]:
    return [cell_text(cell) for cell in table.rows[row_idx].cells]


def clear_cell(cell) -> None:
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.text = ""
    if not cell.paragraphs:
        cell.add_paragraph("")


def set_cell_text(cell, text: str) -> None:
    clear_cell(cell)
    if cell.paragraphs:
        paragraph = cell.paragraphs[0]
        if paragraph.runs:
            paragraph.runs[0].text = str(text)
        else:
            paragraph.add_run(str(text))
    else:
        cell.add_paragraph(str(text))


class ScheduleWordEditor:
    def __init__(self, path: str | Path):
        self.path = Path(path)
        self.doc = Document(str(self.path))
        if not self.doc.tables:
            raise ValueError("Word 檔沒有表格，這版目前無法處理。")
        self.table = self.doc.tables[0]
        self.document_year = self._detect_document_year()

    def save_as(self, output_path: str | Path) -> str:
        output = str(output_path)
        self.doc.save(output)
        return output

    def autofix_template(self, actions: list[ScheduleAction] | None = None) -> list[str]:
        messages: list[str] = []
        relevant_months = self._collect_relevant_months(actions or [])

        for header in self.list_day_headers():
            corrected = self._correct_header_weekday(header.row_idx)
            if corrected:
                messages.append(corrected)

        for month in sorted(relevant_months):
            inserted = self.complete_month(month, mode="range", action_dates=actions or [])
            if inserted:
                messages.append(f"已補齊 {month} 月日期：{', '.join(inserted)}")

        for header in self.list_day_headers():
            corrected = self._correct_header_weekday(header.row_idx)
            if corrected:
                messages.append(corrected)

        return self._dedupe_messages(messages)

    def find_day_row(self, date_text: str) -> int | None:
        target = normalize_date(date_text)
        for idx in range(len(self.table.rows)):
            merged = " | ".join(row_texts(self.table, idx))
            if target in merged and DAY_HEADER_PATTERN.search(merged):
                return idx
        return None

    def find_next_day_row(self, start_row: int) -> int:
        for idx in range(start_row + 1, len(self.table.rows)):
            merged = " | ".join(row_texts(self.table, idx))
            if DAY_HEADER_PATTERN.search(merged):
                return idx
        return len(self.table.rows)

    def add_action(self, action: ScheduleAction) -> str:
        day_row = self.ensure_day_row(action.date)

        entries = self._read_day_entries(day_row)
        normalized_time = normalize_time(action.time)
        replaced = False

        for entry in entries:
            if normalized_time and entry["time"] == normalized_time:
                entry["time"] = normalized_time
                entry["event"] = action.event
                entry["address"] = action.address
                entry["note"] = action.note
                replaced = True
                break

        if not replaced:
            entries.append(
                {
                    "time": normalized_time,
                    "event": action.event,
                    "address": action.address,
                    "note": action.note,
                }
            )

        self._write_day_entries(day_row, entries)
        if replaced:
            return f"已更新 {action.date} {normalized_time or action.event}"
        return f"已加入 {action.date} {normalized_time or action.event}"

    def ensure_day_row(self, date_text: str) -> int:
        normalized_target = normalize_date(date_text)
        existing = self.find_day_row(normalized_target)
        if existing is not None:
            return existing

        target_date = self._date_in_document_year(normalized_target)
        previous_header, next_header = self._find_neighbor_headers(target_date)
        if previous_header is None and next_header is None:
            raise ValueError("Word 表格裡找不到任何日期標題列，無法自動補日期。")

        if previous_header and next_header:
            self._fill_gap_after_header(previous_header, target_date, stop_before=next_header.row_idx)
        elif previous_header:
            self._fill_gap_after_header(previous_header, target_date)
        elif next_header:
            self._fill_gap_before_header(next_header, target_date)

        created = self.find_day_row(normalized_target)
        if created is None:
            raise ValueError(f"已嘗試補上 {normalized_target}，但仍找不到該日期列。")
        return created

    def delete_action(self, action: ScheduleAction) -> bool:
        day_row = self.find_day_row(action.date)
        if day_row is None:
            return False

        if action.action == ActionType.DELETE_DAY:
            next_day_row = self.find_next_day_row(day_row)
            for tr in [self.table.rows[i]._tr for i in range(day_row, next_day_row)]:
                self.table._tbl.remove(tr)
            return True

        target_time = normalize_time(action.time)
        entries = self._read_day_entries(day_row)
        kept = [entry for entry in entries if entry["time"] != target_time]
        if len(kept) == len(entries):
            return False
        self._write_day_entries(day_row, kept)
        return True

    def apply(self, actions: list[ScheduleAction], output_dir: str | Path) -> ApplySummary:
        summary = ApplySummary()
        summary.messages.extend(self.autofix_template(actions))

        for action in actions:
            try:
                if action.action == ActionType.ADD:
                    summary.messages.append(self.add_action(action))
                    summary.applied.append(f"{action.action.value}:{action.date}:{action.time}:{action.event}")
                else:
                    ok = self.delete_action(action)
                    if ok:
                        summary.messages.append(f"已刪除 {action.date} {action.time}".strip())
                        summary.applied.append(f"{action.action.value}:{action.date}:{action.time}")
                    else:
                        summary.messages.append(f"找不到 {action.date} {action.time}".strip())
                        summary.skipped.append(f"{action.action.value}:{action.date}:{action.time}")
            except Exception as exc:
                summary.messages.append(str(exc))
                summary.skipped.append(f"{action.action.value}:{action.date}:{action.time}:{action.event}")

        output_dir = Path(output_dir)
        output_dir.mkdir(parents=True, exist_ok=True)
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        output_path = output_dir / f"{self.path.stem}_line_mvp_{timestamp}{self.path.suffix}"
        summary.output_path = self.save_as(output_path)
        return summary

    def list_day_headers(self) -> list[DayHeader]:
        headers: list[DayHeader] = []
        for row_idx in range(len(self.table.rows)):
            texts = row_texts(self.table, row_idx)
            for text in texts:
                match = DAY_HEADER_PATTERN.search(text)
                if match:
                    headers.append(
                        DayHeader(
                            row_idx=row_idx,
                            month=int(match.group(1)),
                            day=int(match.group(2)),
                            raw_text=text,
                        )
                    )
                    break
        return headers

    def complete_month(
        self,
        month: int,
        *,
        mode: str = "range",
        action_dates: list[ScheduleAction] | None = None,
    ) -> list[str]:
        headers = [header for header in self.list_day_headers() if header.month == month]
        action_days = [
            int(normalize_date(action.date).split("/")[1])
            for action in (action_dates or [])
            if normalize_date(action.date).startswith(f"{month}/")
        ]
        if not headers and not action_days:
            return []

        last_day = monthrange(self.document_year, month)[1]
        existing_days = {header.day for header in headers}

        if mode == "full":
            start_day = 1
            end_day = last_day
        else:
            candidates = list(existing_days) + action_days
            start_day = min(candidates)
            end_day = max(candidates)

        inserted_dates: list[str] = []
        for day_value in range(start_day, end_day + 1):
            if day_value in existing_days:
                continue
            self.insert_clean_day_block(month, day_value)
            inserted_dates.append(f"{month}/{day_value}")
            existing_days.add(day_value)
        return inserted_dates

    def insert_clean_day_block(self, month: int, day_value: int) -> int:
        date_text = f"{month}/{day_value}"
        existing = self.find_day_row(date_text)
        if existing is not None:
            return existing

        target_date = date(self.document_year, month, day_value)
        previous_header, next_header = self._find_neighbor_headers(target_date)

        if previous_header:
            self._fill_gap_after_header(previous_header, target_date, stop_before=next_header.row_idx if next_header else None)
        elif next_header:
            self._fill_gap_before_header(next_header, target_date)
        else:
            raise ValueError("找不到可複製的日期模板。")

        row = self.find_day_row(date_text)
        if row is None:
            raise ValueError(f"插入日期 {date_text} 失敗。")
        return row

    def _read_day_entries(self, day_row: int) -> list[dict[str, str]]:
        entries: list[dict[str, str]] = []
        for row_idx in range(day_row + 1, self.find_next_day_row(day_row)):
            texts = row_texts(self.table, row_idx)
            if DAY_HEADER_PATTERN.search(" | ".join(texts)):
                break
            time_text = normalize_time(texts[1] if len(texts) > 1 else "")
            event = texts[2].strip() if len(texts) > 2 else ""
            address = texts[3].strip() if len(texts) > 3 else ""
            note = texts[4].strip() if len(texts) > 4 else ""
            if not any((time_text, event, address, note)):
                continue
            entries.append(
                {
                    "time": time_text,
                    "event": event,
                    "address": address,
                    "note": note,
                }
            )
        entries.sort(key=lambda item: (time_sort_key(item["time"]), item["event"]))
        return entries

    def _write_day_entries(self, day_row: int, entries: list[dict[str, str]]) -> None:
        next_day_row = self.find_next_day_row(day_row)
        row_indexes = list(range(day_row + 1, next_day_row))
        if not row_indexes:
            raise ValueError("日期區塊內沒有可用列。")

        while len(row_indexes) < len(entries):
            template_row_idx = row_indexes[-1]
            inserted_idx = self._insert_row_before(next_day_row, template_row_idx)
            row_indexes.append(inserted_idx)
            next_day_row += 1

        for idx, entry in zip(row_indexes, entries):
            self._fill_schedule_row(idx, entry)

        for idx in row_indexes[len(entries):]:
            row = self.table.rows[idx]
            for cell_index in range(1, min(5, len(row.cells))):
                set_cell_text(row.cells[cell_index], "")

    def _fill_schedule_row(self, row_idx: int, entry: dict[str, str]) -> None:
        row = self.table.rows[row_idx]
        if len(row.cells) > 1:
            set_cell_text(row.cells[1], entry["time"])
        if len(row.cells) > 2:
            set_cell_text(row.cells[2], entry["event"])
        if len(row.cells) > 3:
            set_cell_text(row.cells[3], entry["address"])
        if len(row.cells) > 4:
            set_cell_text(row.cells[4], entry["note"])

    def _fill_gap_after_header(
        self,
        header: DayHeader,
        target_date: date,
        *,
        stop_before: int | None = None,
    ) -> None:
        anchor_row = self.find_next_day_row(header.row_idx) if stop_before is None else stop_before
        template_row = header.row_idx
        cursor = self._header_date(header) + timedelta(days=1)
        while cursor <= target_date:
            inserted_row = self._clone_day_block(template_row, anchor_row, cursor)
            template_row = inserted_row
            anchor_row = self.find_next_day_row(inserted_row)
            cursor += timedelta(days=1)

    def _fill_gap_before_header(self, header: DayHeader, target_date: date) -> None:
        missing_dates: list[date] = []
        cursor = target_date
        end_date = self._header_date(header)
        while cursor < end_date:
            missing_dates.append(cursor)
            cursor += timedelta(days=1)

        template_row = header.row_idx
        anchor_row = header.row_idx
        for missing_date in reversed(missing_dates):
            inserted_row = self._clone_day_block(template_row, anchor_row, missing_date)
            template_row = inserted_row
            anchor_row = inserted_row

    def _clone_day_block(self, template_day_row: int, insert_before: int, target_date: date) -> int:
        block_size = self.find_next_day_row(template_day_row) - template_day_row
        if block_size <= 0:
            raise ValueError("日期區塊格式異常，無法複製日期模板。")

        template_rows = [deepcopy(self.table.rows[template_day_row + offset]._tr) for offset in range(block_size)]
        inserted_rows: list[int] = []

        if insert_before >= len(self.table.rows):
            for new_tr in template_rows:
                self.table._tbl.append(new_tr)
                inserted_rows.append(self._locate_row(new_tr))
        else:
            anchor = self.table.rows[insert_before]._tr
            previous = anchor
            for index, new_tr in enumerate(template_rows):
                if index == 0:
                    previous.addprevious(new_tr)
                else:
                    previous.addnext(new_tr)
                previous = new_tr
                inserted_rows.append(self._locate_row(new_tr))

        header_row_idx = inserted_rows[0]
        self._set_day_header_text(header_row_idx, self._format_date(target_date))
        for row_idx in inserted_rows[1:]:
            row = self.table.rows[row_idx]
            for cell_index in range(1, min(5, len(row.cells))):
                set_cell_text(row.cells[cell_index], "")
        return header_row_idx

    def _insert_row_before(self, row_idx: int, template_row_idx: int) -> int:
        new_tr = deepcopy(self.table.rows[template_row_idx]._tr)
        if row_idx >= len(self.table.rows):
            self.table._tbl.append(new_tr)
        else:
            self.table.rows[row_idx]._tr.addprevious(new_tr)
        return self._locate_row(new_tr)

    def _find_neighbor_headers(self, target_date: date) -> tuple[DayHeader | None, DayHeader | None]:
        previous_header: DayHeader | None = None
        next_header: DayHeader | None = None
        for header in self.list_day_headers():
            header_date = self._header_date(header)
            if header_date < target_date:
                previous_header = header
                continue
            if header_date > target_date:
                next_header = header
                break
        return previous_header, next_header

    def _header_date(self, header: DayHeader) -> date:
        return date(self.document_year, header.month, header.day)

    def _set_day_header_text(self, row_idx: int, date_text: str) -> None:
        row = self.table.rows[row_idx]
        normalized = normalize_date(date_text)
        header_text = self._build_header_text(normalized)
        for cell in row.cells:
            original = cell_text(cell)
            if DAY_HEADER_PATTERN.search(original):
                weekday_match = HEADER_WITH_WEEKDAY_PATTERN.search(original)
                if weekday_match:
                    updated = HEADER_WITH_WEEKDAY_PATTERN.sub(
                        lambda match: (
                            f"{normalized}"
                            f"{match.group('space')}"
                            f"{match.group('open')}"
                            f"{self._weekday_text(normalized)}"
                            f"{match.group('close')}"
                        ),
                        original,
                        count=1,
                    )
                else:
                    updated = re.sub(r"\d{1,2}/\d{1,2}", normalized, original, count=1)
                set_cell_text(cell, updated)
                return
        target_index = 1 if len(row.cells) > 1 else 0
        set_cell_text(row.cells[target_index], header_text)

    def _correct_header_weekday(self, row_idx: int) -> str:
        row = self.table.rows[row_idx]
        normalized = self._extract_header_date(row_idx)
        if not normalized:
            return ""

        expected_weekday = self._weekday_text(normalized)
        for cell in row.cells:
            original = cell_text(cell)
            match = HEADER_WITH_WEEKDAY_PATTERN.search(original)
            if not match:
                continue
            if match.group("weekday") == expected_weekday:
                return ""
            updated = HEADER_WITH_WEEKDAY_PATTERN.sub(
                lambda current: (
                    f"{normalize_date(normalized)}"
                    f"{current.group('space')}"
                    f"{current.group('open')}"
                    f"{expected_weekday}"
                    f"{current.group('close')}"
                ),
                original,
                count=1,
            )
            set_cell_text(cell, updated)
            return f"已校正星期：{normalize_date(normalized)} -> {expected_weekday}"
        return ""

    def _extract_header_date(self, row_idx: int) -> str:
        for text in row_texts(self.table, row_idx):
            match = DAY_HEADER_PATTERN.search(text)
            if match:
                return f"{int(match.group(1))}/{int(match.group(2))}"
        return ""

    def _locate_row(self, target_tr) -> int:
        for idx in range(len(self.table.rows)):
            if self.table.rows[idx]._tr is target_tr:
                return idx
        raise RuntimeError("找不到剛插入的列。")

    def _collect_relevant_months(self, actions: list[ScheduleAction]) -> set[int]:
        months = {header.month for header in self.list_day_headers()}
        for action in actions:
            normalized = normalize_date(action.date)
            match = re.fullmatch(r"(\d{1,2})/(\d{1,2})", normalized)
            if match:
                months.add(int(match.group(1)))
        return months

    def _detect_document_year(self) -> int:
        texts = [paragraph.text for paragraph in self.doc.paragraphs]
        for row in self.table.rows:
            for cell in row.cells:
                texts.append(cell_text(cell))

        merged = "\n".join(texts)
        ad_match = re.search(r"(20\d{2})年", merged)
        if ad_match:
            return int(ad_match.group(1))

        roc_match = re.search(r"(?<!\d)(1\d{2})年", merged)
        if roc_match:
            return int(roc_match.group(1)) + 1911

        stem_match = re.match(r"^(1\d{2})\d{4}", self.path.stem)
        if stem_match:
            return int(stem_match.group(1)) + 1911

        return datetime.now().year

    def _date_in_document_year(self, date_text: str) -> date:
        normalized = normalize_date(date_text)
        match = re.fullmatch(r"(\d{1,2})/(\d{1,2})", normalized)
        if not match:
            raise ValueError(f"無法辨識日期格式：{date_text}")
        return date(self.document_year, int(match.group(1)), int(match.group(2)))

    def _build_header_text(self, normalized_date: str) -> str:
        return f"{normalized_date} ({self._weekday_text(normalized_date)})"

    def _weekday_text(self, normalized_date: str) -> str:
        target = self._date_in_document_year(normalized_date)
        return WEEKDAY_MAP[target.weekday()]

    @staticmethod
    def _format_date(value: date) -> str:
        return f"{value.month}/{value.day}"

    @staticmethod
    def _dedupe_messages(messages: list[str]) -> list[str]:
        seen: set[str] = set()
        result: list[str] = []
        for message in messages:
            if not message or message in seen:
                continue
            seen.add(message)
            result.append(message)
        return result
