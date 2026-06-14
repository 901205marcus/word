from __future__ import annotations

import base64
import mimetypes
import os
from io import BytesIO
from pathlib import Path

from docx import Document
from openai import OpenAI


class AttachmentReader:
    def __init__(self, base_dir: str | Path):
        self.base_dir = Path(base_dir)

    def extract_text(self, filename: str, content: bytes) -> str:
        suffix = Path(filename).suffix.lower()
        if suffix in {".txt", ".md", ".csv", ".log"}:
            return self._decode_text(content)
        if suffix == ".docx":
            return self._read_docx(content)
        if suffix == ".pdf":
            return self._read_pdf(content)
        if suffix in {".png", ".jpg", ".jpeg", ".webp", ".bmp"}:
            return self._read_image_with_openai(filename, content)
        raise ValueError(f"目前不支援這種附件格式：{suffix or filename}")

    @staticmethod
    def _decode_text(content: bytes) -> str:
        for encoding in ("utf-8", "utf-8-sig", "cp950", "big5"):
            try:
                return content.decode(encoding).strip()
            except UnicodeDecodeError:
                continue
        return content.decode("latin-1", errors="ignore").strip()

    @staticmethod
    def _read_docx(content: bytes) -> str:
        document = Document(BytesIO(content))
        parts: list[str] = []
        for paragraph in document.paragraphs:
            text = paragraph.text.strip()
            if text:
                parts.append(text)
        for table in document.tables:
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if cells:
                    parts.append(" | ".join(cells))
        return "\n".join(parts).strip()

    @staticmethod
    def _read_pdf(content: bytes) -> str:
        try:
            from pypdf import PdfReader
        except ImportError as exc:
            raise ValueError("目前伺服器尚未安裝 PDF 讀取套件。") from exc

        reader = PdfReader(BytesIO(content))
        parts: list[str] = []
        for page in reader.pages:
            text = (page.extract_text() or "").strip()
            if text:
                parts.append(text)
        if not parts:
            raise ValueError("這份 PDF 沒有可擷取文字，可能是掃描檔，之後可再補 OCR。")
        return "\n".join(parts).strip()

    @staticmethod
    def _read_image_with_openai(filename: str, content: bytes) -> str:
        api_key = os.getenv("OPENAI_API_KEY", "")
        if not api_key:
            raise ValueError("要讀取圖片文字，請先設定 OPENAI_API_KEY。")

        mime_type = mimetypes.guess_type(filename)[0] or "image/png"
        data_url = f"data:{mime_type};base64,{base64.b64encode(content).decode('utf-8')}"
        client = OpenAI(api_key=api_key)
        response = client.responses.create(
            model="gpt-4.1-mini",
            input=[
                {
                    "role": "user",
                    "content": [
                        {
                            "type": "input_text",
                            "text": "請精準擷取圖片中的所有文字。只輸出純文字，不要解釋。",
                        },
                        {"type": "input_image", "image_url": data_url},
                    ],
                }
            ],
        )
        return (getattr(response, "output_text", "") or "").strip()
