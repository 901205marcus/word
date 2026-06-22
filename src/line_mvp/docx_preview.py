from __future__ import annotations

import math
from pathlib import Path

from PIL import Image, ImageDraw, ImageFont
from docx import Document


def cell_text(cell) -> str:
    return "\n".join(paragraph.text.strip() for paragraph in cell.paragraphs).strip()


class DocxPreviewRenderer:
    def __init__(self):
        self.font = self._load_font(26)
        self.small_font = self._load_font(22)
        self.max_page_height = 3200

    def render(self, docx_path: str | Path, output_dir: str | Path) -> list[str]:
        document = Document(str(docx_path))
        if not document.tables:
            return []

        rows = [[cell_text(cell) for cell in row.cells] for row in document.tables[0].rows]
        if not rows:
            return []

        output_dir = Path(output_dir)
        output_dir.mkdir(parents=True, exist_ok=True)

        page_width = 2200
        padding = 30
        col_widths = self._build_column_widths(rows[0], page_width - padding * 2)
        rendered_rows = self._measure_rows(rows, col_widths)
        stem = Path(docx_path).stem
        pages = self._paginate_rows(rendered_rows)

        output_paths: list[str] = []
        for page_index, page_rows in enumerate(pages, start=1):
            total_height = padding * 2 + sum(row["height"] for row in page_rows)
            image = Image.new("RGB", (page_width, total_height), "white")
            draw = ImageDraw.Draw(image)

            y = padding
            for row in page_rows:
                x = padding
                fill = "#f7e8d7" if row.get("is_header") else "white"
                for col_index, text_lines in enumerate(row["lines"]):
                    width = col_widths[col_index]
                    draw.rectangle((x, y, x + width, y + row["height"]), outline="#b89a6a", fill=fill)
                    self._draw_multiline(draw, x + 12, y + 10, text_lines, self.small_font, width - 24)
                    x += width
                y += row["height"]

            suffix = "" if len(pages) == 1 else f"_{page_index:02d}"
            target = output_dir / f"{stem}_preview{suffix}.jpg"
            image.save(target, format="JPEG", quality=92)
            output_paths.append(str(target))
        return output_paths

    def _measure_rows(self, rows: list[list[str]], col_widths: list[int]) -> list[dict]:
        measured: list[dict] = []
        for row_index, row in enumerate(rows):
            wrapped_cols: list[list[str]] = []
            height = 52
            for index, text in enumerate(row):
                wrapped = self._wrap_text(text or "", col_widths[index] - 24)
                wrapped_cols.append(wrapped)
                height = max(height, 20 + len(wrapped) * 30)
            measured.append({"lines": wrapped_cols, "height": height, "is_header": row_index == 0})
        return measured

    def _paginate_rows(self, rows: list[dict]) -> list[list[dict]]:
        if not rows:
            return []

        pages: list[list[dict]] = []
        header_row = rows[0]
        current_page: list[dict] = [header_row]
        current_height = header_row["height"]

        for row in rows[1:]:
            next_height = current_height + row["height"]
            if current_page and next_height > self.max_page_height:
                pages.append(current_page)
                current_page = [header_row, row]
                current_height = header_row["height"] + row["height"]
                continue
            current_page.append(row)
            current_height = next_height

        if current_page:
            pages.append(current_page)
        return pages

    def _wrap_text(self, text: str, max_width: int) -> list[str]:
        if not text:
            return [""]
        lines: list[str] = []
        for raw_line in text.splitlines() or [text]:
            current = ""
            for char in raw_line:
                candidate = current + char
                bbox = self.small_font.getbbox(candidate)
                if bbox[2] - bbox[0] > max_width and current:
                    lines.append(current)
                    current = char
                else:
                    current = candidate
            lines.append(current or "")
        return lines

    def _draw_multiline(self, draw: ImageDraw.ImageDraw, x: int, y: int, lines: list[str], font, max_width: int) -> None:
        cursor_y = y
        for line in lines:
            draw.text((x, cursor_y), line, fill="#1f1f1f", font=font)
            cursor_y += 30

    def _build_column_widths(self, header_row: list[str], usable_width: int) -> list[int]:
        if len(header_row) >= 5:
            widths = [120, 160, 420, 360, 740]
            if len(header_row) > 5:
                widths.extend([math.floor((usable_width - sum(widths)) / (len(header_row) - 5))] * (len(header_row) - 5))
            return widths[: len(header_row)]
        base = math.floor(usable_width / max(1, len(header_row)))
        return [base] * len(header_row)

    @staticmethod
    def _load_font(size: int):
        candidates = [
            "C:/Windows/Fonts/msjh.ttc",
            "C:/Windows/Fonts/mingliu.ttc",
            "/usr/share/fonts/truetype/noto/NotoSansCJK-Regular.ttc",
            "/usr/share/fonts/opentype/noto/NotoSansCJK-Regular.ttc",
            "/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf",
        ]
        for candidate in candidates:
            path = Path(candidate)
            if path.exists():
                try:
                    return ImageFont.truetype(str(path), size=size)
                except Exception:
                    continue
        return ImageFont.load_default()
