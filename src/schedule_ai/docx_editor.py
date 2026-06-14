from __future__ import annotations

import re
from calendar import monthrange
from copy import deepcopy
from datetime import date, datetime
from typing import Optional

from docx import Document

from .config import AssistantConfig
from .models import ActionType, ApplyResult, ScheduleAction
from .parser import normalize_date, normalize_time

WEEKDAY_MAP = "一二三四五六日"
DAY_HEADER_PATTERN = re.compile(r"(\d{1,2})/(\d{1,2})[\(（]")


def cell_text(cell) -> str:
    return "\n".join(p.text.strip() for p in cell.paragraphs).strip()


def row_texts(table, row_idx: int) -> list[str]:
    return [cell_text(c) for c in table.rows[row_idx].cells]


def clear_cell(cell) -> None:
    for p in cell.paragraphs:
        for run in p.runs:
            run.text = ""
    if not cell.paragraphs:
        cell.add_paragraph("")


def set_cell_text(cell, text: str) -> None:
    clear_cell(cell)
    text = str(text)
    if cell.paragraphs:
        paragraph = cell.paragraphs[0]
        if paragraph.runs:
            paragraph.runs[0].text = text
        else:
            paragraph.add_run(text)
    else:
        cell.add_paragraph(text)


class ScheduleDocEditor:
    def __init__(self, doc: Document, config: AssistantConfig):
        self.doc = doc
        self.config = config
        if not doc.tables:
            raise ValueError("這個 Word 檔沒有表格")
        self.table = doc.tables[0]

    @classmethod
    def from_path(cls, path: str, config: AssistantConfig) -> "ScheduleDocEditor":
        return cls(Document(path), config)

    def save(self, output_path: str) -> None:
        self.doc.save(output_path)

    def detect_document_year(self) -> int:
        texts = [p.text for p in self.doc.paragraphs]
        for row in self.table.rows:
            for cell in row.cells:
                texts.append(cell_text(cell))

        merged = "\n".join(texts)
        match = re.search(r"(20\d{2})年", merged)
        if match:
            return int(match.group(1))
        if self.config.default_year:
            return self.config.default_year
        return datetime.now().year

    def header_text(self, month: int, day_value: int) -> str:
        target = date(self.detect_document_year(), month, day_value)
        weekday = WEEKDAY_MAP[target.weekday()]
        return f"{month}/{day_value}({weekday})"

    def is_day_header_text(self, text: str) -> bool:
        return bool(DAY_HEADER_PATTERN.search(text.strip()))

    def find_day_row(self, date_keyword: str) -> Optional[int]:
        date_keyword = normalize_date(date_keyword)
        for i in range(len(self.table.rows)):
            merged = " | ".join(row_texts(self.table, i))
            if date_keyword in merged and self.is_day_header_text(merged):
                return i
        return None

    def find_next_day_row(self, start_row: int) -> int:
        for i in range(start_row + 1, len(self.table.rows)):
            merged = " | ".join(row_texts(self.table, i))
            if self.is_day_header_text(merged):
                return i
        return len(self.table.rows)

    def list_day_headers(self) -> list[tuple[int, int, int]]:
        headers: list[tuple[int, int, int]] = []
        for row_idx in range(len(self.table.rows)):
            merged = " | ".join(row_texts(self.table, row_idx))
            match = DAY_HEADER_PATTERN.search(merged)
            if match:
                headers.append((row_idx, int(match.group(1)), int(match.group(2))))
        return headers

    def find_time_row_in_day(self, day_row: int, time_keyword: str) -> Optional[int]:
        end_row = self.find_next_day_row(day_row)
        wanted = normalize_time(time_keyword)
        for i in range(day_row + 1, end_row):
            for text in row_texts(self.table, i):
                if normalize_time(text) == wanted:
                    return i
        return None

    def is_blank_schedule_row(self, row_idx: int) -> bool:
        texts = row_texts(self.table, row_idx)
        if self.is_day_header_text(" | ".join(texts)):
            return False
        joined = "".join(texts).strip()
        if not joined:
            return True
        non_empty_count = sum(1 for text in texts if text.strip())
        return non_empty_count <= 1

    def find_blank_row_in_day(self, day_row: int) -> Optional[int]:
        end_row = self.find_next_day_row(day_row)
        for i in range(day_row + 1, end_row):
            if self.is_blank_schedule_row(i):
                return i
        return None

    def insert_row_before(self, row_idx: int, template_row_idx: Optional[int] = None) -> int:
        if template_row_idx is None:
            template_row_idx = max(0, row_idx - 1)
        new_tr = deepcopy(self.table.rows[template_row_idx]._tr)
        if row_idx >= len(self.table.rows):
            self.table._tbl.append(new_tr)
        else:
            ref_tr = self.table.rows[row_idx]._tr
            ref_tr.addprevious(new_tr)
        for i in range(len(self.table.rows)):
            if self.table.rows[i]._tr is new_tr:
                return i
        raise RuntimeError("插入列失敗")

    def _find_header_template_row(self) -> int:
        headers = self.list_day_headers()
        if not headers:
            raise ValueError("找不到日期標題列樣板")
        return headers[0][0]

    def _find_schedule_template_row(self) -> int:
        headers = self.list_day_headers()
        for header_row, _, _ in headers:
            next_row = header_row + 1
            if next_row < len(self.table.rows) and len(self.table.rows[next_row].cells) > 1:
                return next_row
        raise ValueError("找不到空白行程列樣板")

    def _fill_header_row(self, row_idx: int, month: int, day_value: int) -> None:
        row = self.table.rows[row_idx]
        for cell in row.cells:
            clear_cell(cell)
        target_cell = row.cells[1] if len(row.cells) > 1 else row.cells[0]
        set_cell_text(target_cell, self.header_text(month, day_value))

    def _is_blank_header_candidate(self, row_idx: int) -> bool:
        texts = row_texts(self.table, row_idx)
        if self.is_day_header_text(" | ".join(texts)):
            return False
        joined = "".join(texts).strip()
        if joined:
            return False
        next_row_idx = row_idx + 1
        if next_row_idx >= len(self.table.rows):
            return False
        next_joined = "".join(row_texts(self.table, next_row_idx)).strip()
        return bool(next_joined)

    def _resolve_insert_before_row(self, month: int, day_value: int) -> int:
        for row_idx, row_month, row_day in self.list_day_headers():
            if (row_month, row_day) > (month, day_value):
                return row_idx
        return len(self.table.rows)

    def insert_clean_day_block(self, month: int, day_value: int) -> int:
        date_text = f"{month}/{day_value}"
        existing = self.find_day_row(date_text)
        if existing is not None:
            return existing

        insert_before_row = self._resolve_insert_before_row(month, day_value)
        header_template_row = self._find_header_template_row()
        schedule_template_row = self._find_schedule_template_row()

        header_row_idx = self.insert_row_before(insert_before_row, header_template_row)
        self._fill_header_row(header_row_idx, month, day_value)

        schedule_insert_before = header_row_idx + 1
        schedule_row_idx = self.insert_row_before(schedule_insert_before, schedule_template_row)
        for cell in self.table.rows[schedule_row_idx].cells:
            clear_cell(cell)
        return header_row_idx

    def _fill_gap_with_blank_headers(
        self,
        month: int,
        start_day: int,
        end_day: int,
        start_row: int,
        end_row: int,
    ) -> list[int]:
        missing_days = list(range(start_day, end_day + 1))
        blank_candidates = [
            row_idx
            for row_idx in range(start_row, end_row)
            if self._is_blank_header_candidate(row_idx)
        ]

        filled_days: list[int] = []
        for day_value, row_idx in zip(missing_days, blank_candidates):
            self._fill_header_row(row_idx, month, day_value)
            filled_days.append(day_value)
        return filled_days

    def complete_month(self, month: int, year: Optional[int] = None) -> list[str]:
        year = year or self.detect_document_year()
        last_day = monthrange(year, month)[1]
        inserted_dates: list[str] = []

        headers = [
            (row_idx, row_day)
            for row_idx, row_month, row_day in self.list_day_headers()
            if row_month == month
        ]

        if not headers:
            for day_value in range(1, last_day + 1):
                self.insert_clean_day_block(month, day_value)
                inserted_dates.append(f"{month}/{day_value}")
            return inserted_dates

        first_row, first_day = headers[0]
        for day_value in range(1, first_day):
            self.insert_clean_day_block(month, day_value)
            inserted_dates.append(f"{month}/{day_value}")

        refreshed_headers = [
            (row_idx, row_day)
            for row_idx, row_month, row_day in self.list_day_headers()
            if row_month == month
        ]

        for (current_row, current_day), (next_row, next_day) in zip(refreshed_headers, refreshed_headers[1:]):
            if next_day <= current_day + 1:
                continue
            filled_days = self._fill_gap_with_blank_headers(
                month=month,
                start_day=current_day + 1,
                end_day=next_day - 1,
                start_row=current_row + 1,
                end_row=next_row,
            )
            for day_value in filled_days:
                inserted_dates.append(f"{month}/{day_value}")

        final_headers = [
            (row_idx, row_day)
            for row_idx, row_month, row_day in self.list_day_headers()
            if row_month == month
        ]
        existing_days = {row_day for _, row_day in final_headers}

        for day_value in range(1, last_day + 1):
            if day_value in existing_days:
                continue
            self.insert_clean_day_block(month, day_value)
            inserted_dates.append(f"{month}/{day_value}")
            existing_days.add(day_value)

        inserted_dates = sorted(set(inserted_dates), key=lambda item: int(item.split("/")[1]))
        return inserted_dates

    def add_schedule(self, action: ScheduleAction) -> str:
        day_row = self.insert_clean_day_block(*[int(x) for x in normalize_date(action.date).split("/")])
        time_text = normalize_time(action.time or "")

        if time_text:
            existing_row = self.find_time_row_in_day(day_row, time_text)
            if existing_row is not None:
                row = self.table.rows[existing_row]
                self._fill_schedule_row(row, action)
                return f"已覆蓋 {action.date} {time_text}"

        blank_row = self.find_blank_row_in_day(day_row)
        if blank_row is not None:
            self._fill_schedule_row(self.table.rows[blank_row], action)
            return f"已加入 {action.date}"

        next_day_row = self.find_next_day_row(day_row)
        template_row_idx = self._find_schedule_template_row()
        inserted_idx = self.insert_row_before(next_day_row, template_row_idx)
        row = self.table.rows[inserted_idx]
        for cell in row.cells:
            clear_cell(cell)
        self._fill_schedule_row(row, action)
        return f"已插入 {action.date}"

    def _fill_schedule_row(self, row, action: ScheduleAction) -> None:
        if len(row.cells) > 1:
            set_cell_text(row.cells[1], action.time or "")
        if len(row.cells) > 2:
            set_cell_text(row.cells[2], action.event)
        if len(row.cells) > 3:
            set_cell_text(row.cells[3], action.address or self.config.blank_address_placeholder)
        if len(row.cells) > 4:
            set_cell_text(row.cells[4], action.note or self.config.blank_note_placeholder)

    def delete_schedule(self, action: ScheduleAction) -> bool:
        day_row = self.find_day_row(action.date)
        if day_row is None:
            return False
        time_row = self.find_time_row_in_day(day_row, action.time or "")
        if time_row is None:
            return False
        row = self.table.rows[time_row]
        for index in range(1, min(5, len(row.cells))):
            set_cell_text(row.cells[index], "")
        return True

    def delete_day_block(self, date_text: str) -> bool:
        day_row = self.find_day_row(date_text)
        if day_row is None:
            return False
        next_day_row = self.find_next_day_row(day_row)
        for tr in [self.table.rows[i]._tr for i in range(day_row, next_day_row)]:
            self.table._tbl.remove(tr)
        return True

    def prune_past_days(self, today: date) -> list[str]:
        removed: list[str] = []
        year = self.detect_document_year()
        current_rows = len(self.table.rows) - 1
        while current_rows >= 0:
            texts = row_texts(self.table, current_rows)
            merged = " | ".join(texts)
            match = DAY_HEADER_PATTERN.search(merged)
            if match:
                month = int(match.group(1))
                day_value = int(match.group(2))
                row_date = date(year, month, day_value)
                if row_date < today:
                    removed.append(f"{month}/{day_value}")
                    self.delete_day_block(f"{month}/{day_value}")
            current_rows -= 1
        removed.reverse()
        return removed

    def apply_actions(
        self,
        actions: list[ScheduleAction],
        approved_indexes: Optional[set[int]] = None,
    ) -> ApplyResult:
        result = ApplyResult()
        approved_indexes = approved_indexes or set(range(len(actions)))
        for index, action in enumerate(actions):
            if index not in approved_indexes:
                result.skipped.append(action)
                result.messages.append(f"略過 {action.date} {action.event or action.time}")
                continue

            if action.action == ActionType.ADD:
                result.messages.append(self.add_schedule(action))
                result.applied.append(action)
            elif action.action == ActionType.DELETE:
                ok = self.delete_schedule(action)
                result.messages.append(
                    f"已刪除 {action.date} {action.time}" if ok else f"找不到 {action.date} {action.time}"
                )
                if ok:
                    result.applied.append(action)
                else:
                    result.skipped.append(action)
            elif action.action == ActionType.DELETE_DAY:
                ok = self.delete_day_block(action.date)
                result.messages.append(
                    f"已刪除整天 {action.date}" if ok else f"找不到日期 {action.date}"
                )
                if ok:
                    result.applied.append(action)
                else:
                    result.skipped.append(action)
        return result
